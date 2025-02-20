import re
import gensim.downloader as api
from transformers import pipeline
from docx import Document
from docx.shared import RGBColor
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from bag_of_words import load_bag_of_words
from knowledge_graphs import query_similar_terms

# Load pre-trained Word2Vec model from Gensim (Google News embeddings)
word2vec_model = api.load("word2vec-google-news-300")

# Initialize the transformer model for NER (Named Entity Recognition)
nlp = pipeline("ner", model="dslim/bert-base-NER", tokenizer="dslim/bert-base-NER")

# Load Bag of Words dynamically from the JSON file
bag_of_words = load_bag_of_words()

# Enhanced Regular expressions for detecting specific patterns
patterns = {
    "credit_card": r'\b(?:\d{4}[ -]?){3}\d{4}\b|\b3[47]\d{13}\b',  # Matches standard and AmEx credit card formats
    "phone_number": r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b|\(\d{3}\)\s?\d{3}[-.\s]?\d{4}',
    "ssn": r'\b\d{3}-\d{2}-\d{4}\b',
    "email": r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
    "date": r'\b\d{1,2}/\d{1,2}/\d{2,4}\b|\b\w+\s\d{1,2},\s\d{4}\b|\b\d{1,2}-\d{1,2}-\d{2,4}\b',
    "financial_amount": r'\$\d{1,3}(?:,\d{3})*(?:\.\d{2})?|\b\d{1,3}(?:,\d{3})*\s(?:USD|dollars|cents|GBP)\b',
    "passport_number": r'\b[A-Z0-9]{1,2}\d{7,8}\b',
    "ip_address": r'\b(?:\d{1,3}\.){3}\d{1,3}\b',
    "url": r'https?://(?:www\.)?\S+\b',
    "bank_account": r'\b\d{16}\b|\b(?:\d{4}[ -]?){4}\b',
    "vendor_id": r'\bBA\d{4}-\d{4}-\d{8}\b',
    "project_code": r'\b\d{4}/[A-Z]{4}/\d{5}\b',
    "badge_number": r'\b\d{3}-[A-Z]{2}-\d{4}\b',
    "flight_id": r'\bFLGHT\d{4}\b',
    "zip_code": r'\bZIP:\s?\d{5}\b',
    "user_name": r'^[a-zA-Z][a-zA-Z0-9_]{2,19}$'

}

SIMILARITY_THRESHOLD = 0.75

# List of common words and locations that should not be considered sensitive unless in context
COMMON_WORDS = {"i", "you", "he", "she", "it", "we", "they", "my", "your", "his", "her"}
COMMON_LOCATIONS = {"london", "england", "india", "usa", "canada"}

# Colors for highlighting categories in DOCX
CATEGORY_COLORS = {
    "Financial": RGBColor(255, 0, 0),
    "Personal Info": RGBColor(0, 0, 255),
    "Credential": RGBColor(0, 128, 0),
    "Personal Identity": RGBColor(255, 165, 0),
    "Sensitive Entity (NER)": RGBColor(128, 0, 128),
    "Marked Sensitive": RGBColor(139, 69, 19),
    "Project Code": RGBColor(255, 105, 180),
    "Vendor ID": RGBColor(255, 0, 255),
    "Badge Number": RGBColor(0, 128, 128),
    "Flight ID": RGBColor(128, 0, 255),
}

# Function to convert a word into its Word2Vec embedding
def get_word_vector(word):
    if word in word2vec_model:
        return word2vec_model[word]
    else:
        return np.zeros(300)

# Compute cosine similarity
def compute_similarity(vec1, vec2):
    return cosine_similarity([vec1], [vec2])[0][0]

# Detect words similar to those in the Bag of Words using Word2Vec
def find_similar_words(text, bag_of_words):
    detected_words = []
    words_in_text = set(re.findall(r'\w+', text.lower()))

    for word in words_in_text:
        word_vector = get_word_vector(word)
        if np.any(word_vector):
            for sensitive_word, category in bag_of_words.items():
                sensitive_vector = get_word_vector(sensitive_word)
                if np.any(sensitive_vector):
                    similarity = compute_similarity(word_vector, sensitive_vector)
                    if similarity >= SIMILARITY_THRESHOLD:
                        detected_words.append({
                            "word": word,
                            "category": category,
                            "sensitivity": "Medium"
                        })
                        break

    return detected_words

# Read document content
def read_document(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='latin-1') as file:
            content = file.read()
    return content

# NER model-based entity detection
def find_sensitive_data_with_model(text):
    entities = nlp(text)
    detected_entities = []
    current_entity = ""
    current_type = ""
    last_position = 0

    for entity in entities:
        if entity['entity'].startswith('B-') or entity['start'] > last_position + 1:
            if current_entity:
                if current_type == 'B-LOC' and current_entity.lower() in COMMON_LOCATIONS:
                    continue
                detected_entities.append({
                    "word": current_entity.strip(),
                    "category": "Sensitive Entity (NER)",
                    "type": current_type,
                    "sensitivity": "High"
                })
            current_entity = entity['word']
            current_type = entity['entity']
        else:
            current_entity += " " + entity['word']
        last_position = entity['end']

    if current_entity:
        detected_entities.append({
            "word": current_entity.strip(),
            "category": "Sensitive Entity (NER)",
            "type": current_type,
            "sensitivity": "High"
        })

    return detected_entities

# Detect sensitive patterns
def detect_sensitive_patterns(text):
    detected_patterns = []
    for label, pattern in patterns.items():
        matches = re.findall(pattern, text)
        for match in matches:
            category = label.replace("_", " ").title()
            detected_patterns.append({
                "word": match,
                "category": category,
                "sensitivity": "High" if label in ["credit_card", "ssn", "bank_account", "passport_number", "vendor_id", "project_code", "badge_number"] else "Medium"
            })
    return detected_patterns

# Analyze sentence-level sensitivity for locations with pronoun check
def analyze_sentence_sensitivity(text):
    sentences = re.split(r'[.!?]', text)
    context_sensitive = []

    for sentence in sentences:
        sentence = sentence.strip()
        words = sentence.split()
        if any(word.lower() in COMMON_WORDS for word in words):
            for location in COMMON_LOCATIONS:
                if location in sentence.lower():
                    context_sensitive.append({
                        "word": sentence,
                        "category": "Contextual Sensitivity",
                        "sensitivity": "Low"
                    })
                    break

    return context_sensitive

# Mark anything after "is", "was", "are", or "were" as sensitive if a pronoun is present in the sentence
def mark_after_linking_verbs_sensitive(text):
    sentences = re.split(r'[.!?]', text)
    sensitive_phrases = []

    # Linking verbs to check
    linking_verbs = r'\b(?:is|was|are|were)\b'

    for sentence in sentences:
        sentence = sentence.strip()
        if any(word.lower() in COMMON_WORDS for word in sentence.split()):
            # Check for the presence of any of the linking verbs
            match = re.search(fr'{linking_verbs}\s+(.*)', sentence, re.IGNORECASE)
            if match:
                sensitive_part = match.group(1)
                sensitive_phrases.append({
                    "word": sensitive_part.strip(),
                    "category": "Marked Sensitive",
                    "sensitivity": "High"
                })

    return sensitive_phrases

# Merge detected sensitive information from all approaches
def merge_detections(similar_word_detections, ner_detections, pattern_detections, context_sensitive_sentences, sensitive_after_linking_verbs):
    all_detections = similar_word_detections + ner_detections + pattern_detections + context_sensitive_sentences + sensitive_after_linking_verbs
    unique_detections = {item['word']: item for item in all_detections}
    return list(unique_detections.values())

# Highlight sensitive data in the document
def highlight_sensitive_data(doc_content, detections):
    doc = Document()
    paragraphs = doc_content.split('\n')

    for paragraph in paragraphs:
        para = doc.add_paragraph()
        words = paragraph.split()

        for word in words:
            for detection in detections:
                if detection['word'] in word:
                    run = para.add_run(word + f" ({detection['category']}) ")
                    color = CATEGORY_COLORS.get(detection['category'], RGBColor(255, 0, 0))
                    run.font.color.rgb = color
                    break
            else:
                para.add_run(word + " ")

    return doc

# Save the highlighted document
def save_highlighted_docx(doc, output_file):
    doc.save(output_file)

# Main function to detect sensitive data in the document
def detect_sensitive_data(file_path, output_file):
    doc_content = read_document(file_path)

    similar_words_detections = find_similar_words(doc_content, bag_of_words)
    ner_detections = find_sensitive_data_with_model(doc_content)
    pattern_detections = detect_sensitive_patterns(doc_content)
    context_sensitive_sentences = analyze_sentence_sensitivity(doc_content)
    sensitive_after_linking_verbs = mark_after_linking_verbs_sensitive(doc_content)

    all_detections = merge_detections(similar_words_detections, ner_detections, pattern_detections, context_sensitive_sentences, sensitive_after_linking_verbs)
    doc = highlight_sensitive_data(doc_content, all_detections)

    save_highlighted_docx(doc, output_file)
    print(f"Sensitive data detected, highlighted, and saved to {output_file}")

# Usage
detect_sensitive_data('sample_document.txt', 'highlight_sensitive_data.docx')